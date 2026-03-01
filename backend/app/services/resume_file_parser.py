import io
from fastapi import UploadFile


async def extract_text_from_file(file: UploadFile) -> str:
    """Extract plain text from an uploaded PDF, DOCX, or TXT resume file."""
    content = await file.read()
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        return _extract_from_pdf(content)
    elif filename.endswith(".docx"):
        return _extract_from_docx(content)
    else:
        # Plain text / unknown — decode as UTF-8
        return content.decode("utf-8", errors="ignore")


def _extract_from_pdf(content: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception as exc:
        return f"[PDF parse error: {exc}]"


def _extract_from_docx(content: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also grab table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as exc:
        return f"[DOCX parse error: {exc}]"
